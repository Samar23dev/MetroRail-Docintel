from flask import Flask, request, jsonify, make_response, send_file
from flask_cors import CORS
from pymongo import MongoClient
from bson.objectid import ObjectId
from datetime import datetime
import os
import math
from collections import Counter
from dotenv import load_dotenv
import json
from io import BytesIO
from PIL import Image, ImageOps, ImageEnhance, ImageFilter
from werkzeug.utils import secure_filename
from pathlib import Path

try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    pytesseract = None
    OCR_AVAILABLE = False

# --- New Imports for File Processing & API Calls ---
import requests
import fitz  # PyMuPDF
import docx # python-docx

# NOTE: For Windows users, install Tesseract-OCR separately from:
# https://github.com/UB-Mannheim/tesseract/wiki
# Or: choco install tesseract (if using Chocolatey)

load_dotenv()

app = Flask(__name__)
CORS(app)

# MongoDB Connection
MONGO_URI = os.getenv('MONGO_URI', 'mongodb://localhost:27017/')
DB_NAME = os.getenv('DB_NAME', 'kmrl_docintel')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY') # --- IMPORTANT: Add this to your .env file ---

# === GEMINI MODEL SELECTION ===
# Set GEMINI_MODEL in .env to switch between models:
# - 'flash' (default): Fast, cheaper, good for basic extraction
# - 'pro': Advanced handwriting, charts, diagrams recognition
GEMINI_MODEL = os.getenv('GEMINI_MODEL', 'flash')  # Default to 'flash'
OCR_LANGUAGES = os.getenv('OCR_LANGUAGES', 'eng+mal')
OCR_PSM = os.getenv('OCR_PSM', '6')
OCR_OEM = os.getenv('OCR_OEM', '1')
OCR_CONFIG = os.getenv('OCR_CONFIG', '').strip()
UPLOAD_FOLDER = Path(os.getenv('UPLOAD_FOLDER', Path(__file__).parent / 'uploads'))
UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)

if GEMINI_MODEL.lower() == 'pro':
    GEMINI_API_MODEL = 'gemini-1.5-pro'
    print("⭐ Using GEMINI_1.5_PRO for advanced document analysis (handwriting, charts, diagrams)")
else:
    GEMINI_API_MODEL = 'gemini-1.5-flash'
    print("⚡ Using GEMINI_1.5_FLASH for fast document analysis")

client = MongoClient(MONGO_URI)
db = client[DB_NAME]
documents_collection = db['documents']

# Create indexes
documents_collection.create_index('title')
documents_collection.create_index('department')
documents_collection.create_index('type')
documents_collection.create_index('tags')


def serialize_document(doc):
    """Convert Mongo document to JSON-serializable dict with file URLs."""
    if not doc:
        return None
    doc['_id'] = str(doc['_id'])
    if doc.get('file_path'):
        doc['file_url'] = f"/api/documents/{doc['_id']}/file"
    return doc


# -------------------------
# FILE EXTRACTION FUNCTION
# -------------------------

def extract_text_from_file(file_storage):
    """Extracts raw text from an uploaded file (PDF, DOCX, TXT, DOC, XLS, XLSX, JPG, PNG)."""
    filename = file_storage.filename
    text = ""
    try:
        # Reset file pointer to the beginning
        file_storage.seek(0)
        def preprocess_image_for_ocr(image):
            """Enhance handwritten scans for better OCR results."""
            if image.mode not in ('L', 'LA'):
                image = image.convert('L')
            image = ImageOps.autocontrast(image)
            image = ImageEnhance.Contrast(image).enhance(1.8)
            image = image.filter(ImageFilter.MedianFilter(size=3))
            return image
        def run_ocr(image):
            if not OCR_AVAILABLE:
                raise ImportError("pytesseract not installed")
            config_parts = [f"--psm {OCR_PSM}", f"--oem {OCR_OEM}"]
            if OCR_CONFIG:
                config_parts.append(OCR_CONFIG)
            config_str = " ".join(part.strip() for part in config_parts if part)
            image = preprocess_image_for_ocr(image)
            return pytesseract.image_to_string(
                image,
                lang=OCR_LANGUAGES,
                config=config_str
            )
        
        # PDF files
        if filename.endswith('.pdf'):
            pdf_document = fitz.open(stream=file_storage.read(), filetype="pdf")
            
            # First, try to extract text directly (for text-based PDFs)
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                blocks = page.get_text("blocks")
                blocks.sort(key=lambda b: (b[1], b[0]))
                for b in blocks:
                    text += b[4]
            
            # If no text was extracted, try OCR on PDF images (for scanned PDFs)
            if not text.strip():
                try:
                    print(f"[OCR] No text found in PDF {filename}, attempting OCR on images...")
                    
                    for page_num in range(len(pdf_document)):
                        page = pdf_document.load_page(page_num)
                        # Render page as image at higher resolution for better OCR
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                        img_data = pix.tobytes("ppm")
                        image = Image.open(BytesIO(img_data))
                        
                        # Extract text using OCR
                        page_text = run_ocr(image)
                        text += f"\n--- Page {page_num + 1} (OCR) ---\n{page_text}\n"
                    
                    if text.strip():
                        print(f"[OCR] Successfully extracted text from {filename} using OCR")
                    else:
                        print(f"[OCR] No text could be extracted even with OCR for {filename}")
                
                except ImportError as e:
                    print(f"[WARNING] OCR libraries not available for {filename}: {str(e)}")
                    text = "[Scanned PDF detected but OCR capability not available. Please install Pillow and pytesseract.]"
                except Exception as e:
                    print(f"[ERROR] OCR processing failed for {filename}: {str(e)}")
                    text = f"[Scanned PDF detected but OCR failed: {str(e)}]"
            
            pdf_document.close()
        
        # DOCX files (modern Word format)
        elif filename.endswith('.docx'):
            doc = docx.Document(file_storage)
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            if doc.tables:
                text += "\n\n--- Extracted Tables ---\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        text += row_text + "\n"
                    text += "------------------------\n"
        
        # DOC files (older Word format - extract as text)
        elif filename.endswith('.doc'):
            # For .doc files, we try to extract using docx if possible, otherwise return placeholder
            try:
                doc = docx.Document(file_storage)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except:
                # If docx library can't read it, treat as text
                file_storage.seek(0)
                try:
                    text = file_storage.read().decode('utf-8', errors='ignore')
                except:
                    text = "[Document appears to be in .doc format - please convert to .docx for better extraction]"
        
        # TXT files
        elif filename.endswith('.txt'):
            text = file_storage.read().decode('utf-8')
        
        # XLS/XLSX files (Excel) - extract as text
        elif filename.endswith('.xls') or filename.endswith('.xlsx'):
            try:
                import openpyxl
                from openpyxl import load_workbook
                
                file_storage.seek(0)
                workbook = load_workbook(file_storage)
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"\n--- Sheet: {sheet_name} ---\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        if row_text.strip():
                            text += row_text + "\n"
            except ImportError:
                # openpyxl not installed, extract what we can
                file_storage.seek(0)
                try:
                    text = file_storage.read().decode('utf-8', errors='ignore')
                except:
                    text = "[Excel file detected but cannot be processed without openpyxl library]"
        
        # Image files (JPG, PNG) - OCR or placeholder
        elif filename.endswith('.jpg') or filename.endswith('.jpeg') or filename.endswith('.png'):
            try:
                file_storage.seek(0)
                image = Image.open(file_storage)
                text = run_ocr(image)
                
                if not text.strip():
                    text = "[Image file detected - no readable text extracted. Please ensure image contains text.]"
            except ImportError:
                # OCR libraries not installed
                text = "[Image file detected. OCR capability not available. Please install Pillow and pytesseract to enable text extraction from images.]"
            except Exception as e:
                text = f"[Image file detected but could not be processed: {str(e)}]"
        
        else:
            return None  # Unsupported file type
    
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        return None
    
    return text if text.strip() else "[File processed but no text content found]"

# -------------------------
# [UPDATED] GEMINI AI ANALYSIS FUNCTION - NOW INCLUDES STATUS
# -------------------------

# def analyze_document_with_gemini(text_content):
#     """Uses Gemini API to generate summary, tags, and other metadata, including a suggested status."""
    
#     if not GEMINI_API_KEY:
#         print("GEMINI_API_KEY not found. Returning mock data.")
#         return {
#             "title": "Mock Title (GEMINI_API_KEY not set)",
#             "summary": "This is mock data. Please set your GEMINI_API_KEY in .env to enable AI analysis.",
#             "tags": ["error", "mock-data"],
#             "department": "Unknown",
#             "type": "Unknown",
#             "language": "English",
#             "status": "review",
#             "tables_data": [],
#             "figures_data": [],
#             "charts": []
#         }

#     # Truncate text intelligently to avoid exceeding API limits
#     # Gemini 1.5-pro context window is ~1M tokens, but we truncate to be safe
#     # Average English word ≈ 1.3 tokens, so ~12000 words ≈ 15600 tokens
#     truncated_text = text_content[:12000]

#     api_url = f"https://generativelanguage.googleapis.com/v1beta/models/models/gemini-1.5-pro:generateContent?key={GEMINI_API_KEY}"

#     # --- DEFINE RESPONSE SCHEMA FOR JSON EXTRACTION ---
#     response_schema = {
#         "type": "OBJECT",
#         "properties": {
#             "title": {"type": "STRING"},
#             "summary": {"type": "STRING"},
#             "tags": {
#                 "type": "ARRAY",
#                 "items": {"type": "STRING"}
#             },
#             "department": {"type": "STRING"},
#             "type": {"type": "STRING"},
#             "status": {
#                 "type": "STRING",
#                 "enum": ["urgent", "approved", "review"]
#             },
#             "tables_data": { 
#                 "type": "ARRAY",
#                 "items": {
#                     "type": "OBJECT", 
#                     "properties": {
#                         "caption": {"type": "STRING"},
#                         "data": {
#                             "type": "ARRAY",
#                             "items": {
#                                 "type": "ARRAY",
#                                 "items": {"type": "STRING"}
#                             }
#                         }
#                     },
#                     "required": ["caption", "data"]
#                 }
#             },
#             "figures_data": {
#                 "type": "ARRAY",
#                 "items": {
#                     "type": "OBJECT",
#                     "properties": {
#                         "description": {"type": "STRING"},
#                         "values": {
#                             "type": "ARRAY",
#                             "items": {"type": "STRING"}
#                         },
#                         "type": {
#                             "type": "STRING",
#                             "enum": ["number", "percentage", "currency", "metric", "ratio"]
#                         }
#                     },
#                     "required": ["description", "values", "type"]
#                 }
#             },
#             "charts": {
#                 "type": "ARRAY",
#                 "items": {
#                     "type": "OBJECT",
#                     "properties": {
#                         "title": {"type": "STRING"},
#                         "description": {"type": "STRING"},
#                         "chart_type": {
#                             "type": "STRING",
#                             "enum": ["bar", "line", "pie", "area", "scatter"]
#                         },
#                         "data_points": {
#                             "type": "ARRAY",
#                             "items": {
#                                 "type": "OBJECT",
#                                 "properties": {
#                                     "label": {"type": "STRING"},
#                                     "value": {"type": "STRING"}
#                                 }
#                             }
#                         }
#                     },
#                     "required": ["title", "chart_type", "data_points"]
#                 }
#             }
#         },
#         "required": ["title", "summary", "tags", "department", "type", "status", "tables_data", "figures_data", "charts"]
#     }

#     # --- ADVANCED PROMPT FOR PRO MODEL (with handwriting & chart support) ---
#     if GEMINI_MODEL.lower() == 'pro':
#         system_prompt = (
#             "You are an expert document analyst for KMRL (Kochi Metro Rail Limited). "
#             "Analyze this document with advanced capabilities including: "
#             "1. Handwritten text recognition (extract and interpret) "
#             "2. Chart/diagram analysis (type, data points, trends) "
#             "3. Table structure understanding (preserve formatting) "
#             "4. Complex multi-format documents (text, images, tables, charts combined) "
#             "5. Mathematical expressions and technical drawings "
#             "\n"
#             "Return comprehensive JSON with: title, summary (3-4 sentences), tags, department, type, status. "
#             "Status: 'urgent' (critical/incident/deadline), 'approved' (routine/finalized), 'review' (unclear/pending). "
#             "Extract ALL content: tables (with structure), figures, charts, handwritten notes, diagrams. "
#             "For charts: identify type, extract data points, detect trends. "
#             "For handwriting: transcribe text, note legibility issues. "
#             "Return empty arrays only if genuinely absent."
#         )
#     else:
#         # BASIC PROMPT FOR FLASH MODEL (faster, cheaper)
#         system_prompt = (
#             "Analyze document and return JSON with: title, summary (2-3 sentences), tags, department, type, status. "
#             "Status: 'urgent' (critical/incident), 'approved' (routine/finalized), 'review' (unclear). "
#             "Extract: tables (caption + data), figures (description/values/type), "
#             "charts (title/type with data_points). Return empty arrays if none found."
#         )

#     payload = {
#         "systemInstruction": {"parts": [{"text": system_prompt}]},
#         "contents": [{
#             "parts": [{"text": f"Analyze this document: {truncated_text}"}]
#         }],
#         "generationConfig": {
#             "responseMimeType": "application/json",
#             "responseSchema": response_schema
#         }
#     }

#     try:
#         response = requests.post(api_url, json=payload, headers={'Content-Type': 'application/json'})
#         response.raise_for_status() # Raise an error for bad status codes
        
#         result = response.json()
        
#         generated_json_text = result.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
        
#         processed_data = json.loads(generated_json_text)

#         print("--- GEMINI ANALYSIS RESULT ---")
#         print(json.dumps(processed_data, indent=2))
#         print("------------------------------")
        
#         processed_data['language'] = "English" # Default
        
#         return processed_data
        
#     except Exception as e:
#         print(f"Error calling Gemini API: {str(e)}")
#         # Fallback in case of API error
#         return {
#             "title": "Error During Analysis",
#             "summary": f"An error occurred while analyzing the document: {str(e)}",
#             "tags": ["error"],
#             "department": "Unknown",
#             "type": "Unknown",
#             "language": "English",
#             "status": "review",
#             "tables_data": [],
#             "figures_data": [],
#             "charts": []
#         }

import os
import json
from google import genai
from google.genai import types # Used for GenerateContentConfig

# IMPORTANT: Ensure GEMINI_API_KEY is available as an environment variable
# The genai.Client() will automatically look for it.
# You can define a placeholder for this example if needed, but in reality, 
# you should let the client handle environment variable lookup.
# client = genai.Client()
# For demonstration purposes, we'll keep the key checking logic:

def analyze_document_with_gemini(text_content, source_bytes=None, filename=None, mime_type=None, model_type='pro'):
    """
    Uses the Google Gen AI SDK to generate structured document metadata.
    
    :param text_content: The document text to analyze.
    :param model_type: 'pro' or 'flash'.
    :"""
    
    # 1. API KEY CHECK (Using os.getenv for environment variable)
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        print("GEMINI_API_KEY not found. Returning mock data.")
        # --- (Your Mock Data logic remains the same) ---
        return { 
            "title": "Mock Title (GEMINI_API_KEY not set)",
            # ... rest of mock data ... 
            "summary": "This is mock data. Please set your GEMINI_API_KEY in .env to enable AI analysis.",
            "tags": ["error", "mock-data"],
            "department": "Unknown",
            "type": "Unknown",
            "language": "English",
            "status": "review",
            "tables_data": [],
            "figures_data": [],
            "charts": []
        }

    # Initialize Client - SDK handles API key lookup and URL construction
    try:
        client = genai.Client()
    except Exception as e:
        # Handle client initialization error (e.g., if key format is wrong)
        print(f"Error initializing Gemini client: {e}")
        return {
            "title": "Client Initialization Error",
            "summary": "Could not initialize the Gemini client. Check API key format.",
            "tags": ["error", "client-fail"],
            "status": "urgent",
            "tables_data": [], "figures_data": [], "charts": []
        }
        
    # 2. TRUNCATION AND MODEL SELECTION
    # Using the more powerful and current model alias: gemini-2.5-pro
    # For Flash: gemini-2.5-flash
    MODEL_ID = "gemini-2.5-pro" if model_type.lower() == 'pro' else "gemini-2.5-flash"

    # Truncate text (same logic as before)
    truncated_text = (text_content or "")[:12000]

    # 3. DEFINE RESPONSE SCHEMA (Identical to your provided schema)
    response_schema = {
        # ... (response_schema definition from your original code) ...
        # (Putting the schema here for completeness, though it's long)
        "type": "OBJECT",
        "properties": {
            "title": {"type": "STRING"},
            "summary": {"type": "STRING"},
            "tags": {"type": "ARRAY", "items": {"type": "STRING"}},
            "department": {"type": "STRING"},
            "type": {"type": "STRING"},
            "status": {"type": "STRING", "enum": ["urgent", "approved", "review"]},
            "tables_data": { 
                "type": "ARRAY",
                "items": {
                    "type": "OBJECT", 
                    "properties": {
                        "caption": {"type": "STRING"},
                        "data": {"type": "ARRAY", "items": {"type": "ARRAY", "items": {"type": "STRING"}}},
                        "page_number": {"type": "INTEGER"}
                    },
                    "required": ["caption", "data"]
                }
            },
            "figures_data": {
                "type": "ARRAY",
                "items": {
                    "type": "OBJECT",
                    "properties": {
                        "description": {"type": "STRING"},
                        "values": {"type": "ARRAY", "items": {"type": "STRING"}},
                        "type": {"type": "STRING", "enum": ["number", "percentage", "currency", "metric", "ratio"]}
                    },
                    "required": ["description", "values", "type"]
                }
            },
            "charts": {
                "type": "ARRAY",
                "items": {
                    "type": "OBJECT",
                    "properties": {
                        "title": {"type": "STRING"},
                        "description": {"type": "STRING"},
                        "chart_type": {"type": "STRING", "enum": ["bar", "line", "pie", "area", "scatter"]},
                        "data_points": {
                            "type": "ARRAY",
                            "items": {"type": "OBJECT", "properties": {"label": {"type": "STRING"}, "value": {"type": "STRING"}}}
                        },
                        "page_number": {"type": "INTEGER"}
                    },
                    "required": ["title", "chart_type", "data_points"]
                }
            }
        },
        "required": ["title", "summary", "tags", "department", "type", "status", "tables_data", "figures_data", "charts"]
    }

    # 4. SYSTEM PROMPT DEFINITION (Identical to your original logic)
    if model_type.lower() == 'pro':
        system_prompt = (
            "You are an expert document analyst for KMRL (Kochi Metro Rail Limited). "
            "Analyze this document with advanced capabilities including: Handwritten text recognition (extract and interpret), Chart/diagram analysis (type, data points, trends), Table structure understanding (preserve formatting), Complex multi-format documents (text, images, tables, charts combined), and Mathematical expressions and technical drawings. "
            "Return comprehensive JSON with: title, summary (3-4 sentences), tags, department, type, status. "
            "Status: 'urgent' (critical/incident/deadline), 'approved' (routine/finalized), 'review' (unclear/pending). "
            "Extract ALL content: tables (with structure), figures, charts, handwritten notes, diagrams. "
            "For charts: identify type, extract data points, detect trends. "
            "For handwriting: transcribe text, note legibility issues. "
            "IMPORTANT: For each table and chart, include the 'page_number' field indicating which page (1-indexed) the table/chart appears on in the PDF. "
            "If the document mentions page numbers in the text (e.g., 'Page 5', 'on page 3'), use that information. "
            "If page numbers are not explicitly mentioned, estimate based on document structure and content position. "
            "Return empty arrays only if genuinely absent."
        )
    else:
        # BASIC PROMPT FOR FLASH MODEL (faster, cheaper)
        system_prompt = (
            "Analyze document and return JSON with: title, summary (2-3 sentences), tags, department, type, status. "
            "Status: 'urgent' (critical/incident), 'approved' (routine/finalized), 'review' (unclear). "
            "Extract: tables (caption + data + page_number), figures (description/values/type), "
            "charts (title/type with data_points + page_number). "
            "For each table and chart, include 'page_number' (1-indexed) indicating which page it appears on. "
            "Return empty arrays if none found."
        )

    # 5. Build content parts (original file + OCR text)
    content_parts = []

    if source_bytes:
        try:
            file_part = types.Part.from_bytes(
                data=source_bytes,
                mime_type=mime_type or 'application/octet-stream'
            )
            content_parts.append(file_part)
        except Exception as file_err:
            print(f"Error attaching source file to Gemini request: {file_err}")

    if truncated_text:
        content_parts.append(
            types.Part.from_text(
                text="OCR/Extracted text (may include noise but helps with search context):\n"
                f"{truncated_text}"
            )
        )

    if not content_parts:
        content_parts.append(types.Part.from_text(text="No readable content extracted."))

    user_content = types.Content(
        role="user",
        parts=content_parts
    )

    # 6. GENERATION CONFIGURATION (Using SDK's config object)
    config = types.GenerateContentConfig(
        system_instruction=system_prompt, # System prompt is passed in config
        response_mime_type="application/json",
        response_schema=response_schema
    )

    # 7. MAKE THE API CALL
    try:
        response = client.models.generate_content(
            model=MODEL_ID,
            contents=[user_content],
            config=config,
        )
        
        # The response.text is guaranteed to be a valid JSON string due to the config
        processed_data = json.loads(response.text)

        print("--- GEMINI ANALYSIS RESULT (SDK) ---")
        print(json.dumps(processed_data, indent=2))
        print("------------------------------------")
        
        # Add language default
        processed_data['language'] = "English" 
        
        return processed_data
        
    except Exception as e:
        print(f"Error calling Gemini API (SDK): {str(e)}")
        # --- (Your Fallback logic remains the same) ---
        return {
            "title": "Error During Analysis (SDK)",
            "summary": f"An error occurred while analyzing the document using the SDK: {str(e)}",
            "tags": ["error"],
            "department": "Unknown",
            "type": "Unknown",
            "language": "English",
            "status": "review",
            "tables_data": [],
            "figures_data": [],
            "charts": []
        }

# Remember to install the required libraries:
# pip install google-genai
# -------------------------
# CHART NORMALIZATION HELPERS
# -------------------------

def _parse_numeric_value(value):
    """Convert various numeric string formats into floats."""
    if value is None:
        return None

    if isinstance(value, (int, float)):
        return float(value)

    value_str = str(value).strip()
    if not value_str:
        return None

    # Remove common non-numeric prefixes/suffixes (currency, percentage, units)
    value_str = value_str.replace(',', '')
    if value_str.endswith('%'):
        value_str = value_str[:-1]
    # Keep digits, decimal points, and minus signs
    cleaned = ''.join(ch for ch in value_str if ch.isdigit() or ch in '.-')

    if cleaned in {'', '-', '.', '-.', '.-'}:
        return None

    try:
        return float(cleaned)
    except ValueError:
        return None


def _normalize_chart_points(chart):
    """Ensure chart data only contains numeric values so the UI can render accurate visuals."""
    normalized_points = []
    for idx, point in enumerate(chart.get('data_points', [])):
        label = str(point.get('label') or f"Item {idx + 1}").strip()
        numeric_value = _parse_numeric_value(point.get('value'))
        if numeric_value is None:
            continue
        normalized_points.append({
            "label": label,
            "value": numeric_value
        })
    return normalized_points


def _derive_charts_from_tables(tables):
    """
    Build deterministic chart data from extracted tables when AI-provided charts are missing or unreliable.
    Each numeric column (after the first label column) becomes a bar chart.
    """
    derived_charts = []
    tables = tables or []

    for table_index, table in enumerate(tables):
        rows = table.get('data') or []
        if len(rows) < 2:
            continue

        headers = rows[0]
        body_rows = rows[1:]
        if not headers or len(headers) < 2:
            continue

        caption = table.get('caption') or f"Table {table_index + 1}"
        label_header = headers[0] or "Category"

        for col_idx in range(1, len(headers)):
            column_header = headers[col_idx] or f"Column {col_idx + 1}"
            data_points = []

            for row_idx, row in enumerate(body_rows):
                if not row:
                    continue

                label = str(row[0]).strip() if len(row) > 0 and row[0] is not None else f"{label_header} {row_idx + 1}"
                raw_value = row[col_idx] if len(row) > col_idx else None
                numeric_value = _parse_numeric_value(raw_value)

                if label and numeric_value is not None:
                    data_points.append({
                        "label": label,
                        "value": numeric_value
                    })

            if len(data_points) >= 2:
                derived_charts.append({
                    "title": f"{caption} - {column_header}",
                    "description": f"Auto-generated from '{caption}' using '{column_header}' values.",
                    "chart_type": "bar",
                    "data_points": data_points
                })

    return derived_charts


def normalize_charts(processed_data):
    """
    Validate Gemini chart output and fall back to deterministic charts derived from tables if needed.
    Returns a list (possibly empty) that only contains numeric-friendly data so the frontend renders correct charts.
    """
    charts = processed_data.get('charts') or []
    valid_charts = []

    for chart in charts:
        normalized_points = _normalize_chart_points(chart)
        if len(normalized_points) < 2:
            continue

        valid_charts.append({
            "title": chart.get('title') or "Document Chart",
            "description": chart.get('description', ''),
            "chart_type": chart.get('chart_type', 'bar'),
            "data_points": normalized_points
        })

    if valid_charts:
        return valid_charts

    return _derive_charts_from_tables(processed_data.get('tables_data'))

# -------------------------
# SEMANTIC SEARCH (IMPROVED WITH NLP)
# -------------------------

# Common stopwords to exclude from search (improves relevance)
STOPWORDS = {
    'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
    'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during',
    'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
    'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might',
    'this', 'that', 'these', 'those', 'it', 'its', 'they', 'them', 'you', 'i'
}

# Synonym mapping for better semantic matching
SYNONYMS = {
    'safety': ['secure', 'protection', 'security', 'safe'],
    'maintenance': ['repair', 'servicing', 'upkeep', 'maintain'],
    'invoice': ['bill', 'receipt', 'payment', 'charge'],
    'report': ['summary', 'analysis', 'document'],
    'policy': ['rule', 'regulation', 'guideline', 'directive'],
    'urgent': ['critical', 'immediate', 'emergency', 'priority'],
    'engineering': ['technical', 'construction', 'design'],
    'operation': ['operational', 'running', 'execution'],
}

def preprocess_text(text):
    """Preprocess text: lowercase, remove punctuation, filter stopwords."""
    import string
    text = text.lower()
    text = text.translate(str.maketrans('', '', string.punctuation))
    words = text.split()
    words = [w for w in words if w not in STOPWORDS and len(w) > 2]
    return words

def expand_query_with_synonyms(words):
    """Expand query with synonyms for better matching"""
    expanded = list(words)
    for word in words:
        if word in SYNONYMS:
            expanded.extend(SYNONYMS[word])
    return expanded

def text_to_tfidf_vector(text):
    """Convert text to TF-IDF-like vector"""
    words = preprocess_text(text)
    vector = Counter(words)
    boosted_vector = {}
    for word, count in vector.items():
        weight = count * (len(word) / 5.0)
        boosted_vector[word] = weight
    return boosted_vector

def cosine_similarity(vec1, vec2):
    """Enhanced cosine similarity"""
    if not vec1 or not vec2:
        return 0.0
    intersection = set(vec1.keys()) & set(vec2.keys())
    if not intersection:
        return 0.0
    numerator = sum(vec1[x] * vec2[x] for x in intersection)
    sum1 = sum(v * v for v in vec1.values())
    sum2 = sum(v * v for v in vec2.values())
    if sum1 == 0 or sum2 == 0:
        return 0.0
    return numerator / math.sqrt(sum1 * sum2)

def semantic_search(query, collection, top_k=10):
    """Improved semantic search with NLP enhancements"""
    query_words = preprocess_text(query)
    expanded_query_words = expand_query_with_synonyms(query_words)
    query_vec = text_to_tfidf_vector(" ".join(expanded_query_words))
    
    print(f"[SEARCH] Query: '{query}'")
    print(f"[SEARCH] Processed: {query_words}")
    
    results = []
    for doc in collection.find({}):
        full_text = (
            (doc.get("title") or "") + " " +
            (doc.get("summary") or "") + " " +
            (doc.get("department") or "") + " " +
            (doc.get("type") or "") + " " +
            " ".join(doc.get("tags") or [])
        )
        
        doc_vec = text_to_tfidf_vector(full_text)
        score = cosine_similarity(query_vec, doc_vec)
        
        # Boost score for metadata matches
        metadata_boost = 0.0
        query_lower = query.lower()
        if doc.get("department") and doc["department"].lower() in query_lower:
            metadata_boost += 0.3
        if doc.get("type") and doc["type"].lower() in query_lower:
            metadata_boost += 0.3
        for tag in doc.get("tags", []):
            if tag.lower() in query_lower:
                metadata_boost += 0.1
        
        final_score = min(score + metadata_boost, 1.0)
        
        if final_score > 0.05:
            doc["similarity"] = round(final_score, 4)
            doc["_score"] = round(final_score, 4)
            results.append(doc)

    results.sort(key=lambda x: x["similarity"], reverse=True)

    serialized_results = []
    for r in results:
        serialized = serialize_document(r)
        if serialized:
            serialized['similarity'] = r.get('similarity')
            serialized['_score'] = r.get('_score')
            serialized_results.append(serialized)

    return serialized_results[:top_k]


# -------------------------
#       ROUTES
# -------------------------

@app.route('/api/documents', methods=['GET'])
def get_documents():
    try:
        search_query = request.args.get('search', '')
        department = request.args.get('department', '')
        doc_type = request.args.get('type', '')
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10))
        
        query_filter = {}

        if search_query:
            query_filter['$or'] = [
                {'title': {'$regex': search_query, '$options': 'i'}},
                {'summary': {'$regex': search_query, '$options': 'i'}},
                {'tags': {'$regex': search_query, '$options': 'i'}},
                {'content': {'$regex': search_query, '$options': 'i'}}
            ]
        
        if department and department != 'all':
            query_filter['department'] = department
        
        if doc_type and doc_type != 'all-types':
            formatted_type = doc_type.replace('-', ' ').title()
            query_filter['type'] = formatted_type
        
        total = documents_collection.count_documents(query_filter)
        
        skip = (page - 1) * limit
        documents = list(documents_collection.find(query_filter)
                                   .sort('date', -1)
                                   .skip(skip)
                                   .limit(limit))
        documents = [serialize_document(doc) for doc in documents]
        
        return jsonify({
            'documents': documents,
            'pagination': {
                'total': total,
                'page': page,
                'limit': limit,
                'pages': (total + limit - 1) // limit
            }
        })
    except Exception as e:
        print(f"Error in get_documents: {str(e)}")
        return jsonify({'error': str(e)}), 500


# --- [UPDATED] UPLOAD ROUTE - Status is now AI-assigned ---
@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # 1. Extract text from the file
        print(f"Processing file: {file.filename}")
        file.stream.seek(0)
        text_content = extract_text_from_file(file)
        
        if text_content is None:
            return jsonify({'error': 'Unsupported file type or error reading file'}), 400
        
        if not text_content.strip():
            return jsonify({'error': 'File appears to be empty'}), 400
            
        print(f"Extracted {len(text_content)} characters.")

        # 2. Analyze text with Gemini (includes status now)
        file.stream.seek(0)
        original_bytes = file.read()

        # Save original file for future viewing
        original_filename = secure_filename(file.filename) or f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.bin"
        unique_suffix = datetime.now().strftime('%Y%m%d%H%M%S%f')
        stored_filename = f"{unique_suffix}_{original_filename}"
        stored_path = UPLOAD_FOLDER / stored_filename
        with open(stored_path, 'wb') as stored_file:
            stored_file.write(original_bytes)

        print("Analyzing document with Gemini AI...")
        processed_data = analyze_document_with_gemini(
            text_content,
            source_bytes=original_bytes,
            filename=file.filename,
            mime_type=file.mimetype
        )
        print("Analysis complete.")

        # 3a. Ensure chart data is reliable before saving
        processed_data['charts'] = normalize_charts(processed_data)

        # 3. Add remaining data (status is already in processed_data)
        processed_data['content'] = text_content # Store the full text
        processed_data['date'] = datetime.now()
        # processed_data['status'] = 'review' # <-- REMOVED: Status is now AI-assigned
        processed_data['source'] = 'uploaded'
        processed_data['starred'] = False
        processed_data['file_name'] = file.filename
        processed_data['file_mime'] = file.mimetype
        processed_data['file_path'] = str(stored_path)
        
        # 4. Insert into database
        result = documents_collection.insert_one(processed_data)
        processed_data['_id'] = str(result.inserted_id)
        processed_data = serialize_document(processed_data)
        
        print(f"Successfully added document {result.inserted_id} to database.")
        return jsonify(processed_data), 201
        
    except Exception as e:
        print(f"Error in upload_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    try:
        result = documents_collection.delete_one({'_id': ObjectId(doc_id)})
        if result.deleted_count == 0:
            return jsonify({'error': 'Document not found'}), 404
        return jsonify({'message': 'Document deleted successfully'}), 200
    except Exception as e:
        print(f"Error in delete_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/search/semantic', methods=['POST'])
def search_semantic_route():
    try:
        data = request.get_json()
        query = data.get('query', '')

        if not query:
            return jsonify({'error': 'Query is required'}), 400

        results = semantic_search(query, documents_collection)
        
        return jsonify({'results': results})
    except Exception as e:
        print(f"Error in semantic search: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>', methods=['GET'])
def get_document(doc_id):
    try:
        doc = documents_collection.find_one({'_id': ObjectId(doc_id)})
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        
        return jsonify(serialize_document(doc))
    except Exception as e:
        print(f"Error in get_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>', methods=['PUT'])
def update_document(doc_id):
    try:
        data = request.get_json()
        update_data = {}

        if 'status' in data:
            update_data['status'] = data['status']
        if 'tags' in data:
            update_data['tags'] = data['tags']
        if 'starred' in data: 
            update_data['starred'] = data['starred'] 
        
        if not update_data: 
            return jsonify({'error': 'No update fields provided'}), 400

        result = documents_collection.update_one(
            {'_id': ObjectId(doc_id)},
            {'$set': update_data}
        )
        
        if result.matched_count == 0:
            return jsonify({'error': 'Document not found'}), 404
        
        return jsonify({'message': 'Document updated successfully'}), 200
    except Exception as e:
        print(f"Error in update_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>/file', methods=['GET'])
def download_original_file(doc_id):
    """Stream the originally uploaded file for inline viewing."""
    try:
        doc = documents_collection.find_one(
            {'_id': ObjectId(doc_id)},
            {'file_path': 1, 'file_name': 1, 'file_mime': 1}
        )
        if not doc or not doc.get('file_path'):
            return jsonify({'error': 'File not found'}), 404
        
        file_path = Path(doc['file_path'])
        if not file_path.exists():
            return jsonify({'error': 'File missing on server'}), 404
        
        return send_file(
            file_path,
            mimetype=doc.get('file_mime') or 'application/octet-stream',
            as_attachment=False,
            download_name=doc.get('file_name') or file_path.name
        )
    except Exception as e:
        print(f"Error in download_original_file: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/stats', methods=['GET'])
def get_stats():
    try:
        total_docs = documents_collection.count_documents({})
        urgent_docs = documents_collection.count_documents({'status': 'urgent'})
        today_docs = documents_collection.count_documents({
            'date': {'$gte': datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)}
        })
        
        return jsonify({
            'total_documents': total_docs,
            'urgent_items': urgent_docs,
            'documents_today': today_docs
        })
    except Exception as e:
        print(f"Error in get_stats: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/stats', methods=['GET'])
def get_dashboard_stats():
    """Get comprehensive dashboard statistics"""
    try:
        total_docs = documents_collection.count_documents({})
        urgent_docs = documents_collection.count_documents({'status': 'urgent'})
        today_docs = documents_collection.count_documents({
            'date': {'$gte': datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)}
        })
        approved_docs = documents_collection.count_documents({'status': 'approved'})
        review_docs = documents_collection.count_documents({'status': 'review'})
        starred_docs = documents_collection.count_documents({'starred': True})
        
        # Calculate average processing time (mock for now)
        avg_processing_time = 2.3
        
        return jsonify({
            'total_documents': total_docs,
            'urgent_items': urgent_docs,
            'documents_today': today_docs,
            'approved_documents': approved_docs,
            'review_documents': review_docs,
            'starred_documents': starred_docs,
            'avg_processing_time': avg_processing_time,
            'processing_efficiency': 94.2
        })
    except Exception as e:
        print(f"Error in get_dashboard_stats: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/departments', methods=['GET'])
def get_department_stats():
    """Get document count and urgent items per department"""
    try:
        departments_list = [
            "Operations", "Engineering", "Safety", "Procurement",
            "Human Resources", "Finance", "Environment"
        ]
        
        dept_stats = []
        for dept in departments_list:
            total = documents_collection.count_documents({'department': dept})
            urgent = documents_collection.count_documents({
                'department': dept,
                'status': 'urgent'
            })
            approved = documents_collection.count_documents({
                'department': dept,
                'status': 'approved'
            })
            
            dept_stats.append({
                'name': dept,
                'total_docs': total,
                'urgent_items': urgent,
                'approved_items': approved,
                'pending_items': total - urgent - approved
            })
        
        return jsonify({'departments': dept_stats})
    except Exception as e:
        print(f"Error in get_department_stats: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/tags', methods=['GET'])
def get_tags_stats():
    """Get most common tags with frequencies"""
    try:
        all_tags = []
        for doc in documents_collection.find({}, {'tags': 1}):
            if doc.get('tags'):
                all_tags.extend(doc['tags'])
        
        tag_counts = Counter(all_tags)
        top_tags = tag_counts.most_common(10)
        
        tags_data = [
            {'tag': tag, 'count': count}
            for tag, count in top_tags
        ]
        
        return jsonify({'tags': tags_data})
    except Exception as e:
        print(f"Error in get_tags_stats: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/document-types', methods=['GET'])
def get_document_types_stats():
    """Get document count by type"""
    try:
        pipeline = [
            {'$group': {'_id': '$type', 'count': {'$sum': 1}}},
            {'$sort': {'count': -1}},
            {'$limit': 10}
        ]
        
        doc_types = list(documents_collection.aggregate(pipeline))
        
        types_data = [
            {'type': item['_id'] or 'Unknown', 'count': item['count']}
            for item in doc_types
        ]
        
        return jsonify({'document_types': types_data})
    except Exception as e:
        print(f"Error in get_document_types_stats: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/status-distribution', methods=['GET'])
def get_status_distribution():
    """Get document distribution by status"""
    try:
        statuses = ['urgent', 'review', 'approved']
        
        status_data = []
        for status in statuses:
            count = documents_collection.count_documents({'status': status})
            status_data.append({
                'status': status,
                'count': count,
                'percentage': 0
            })
        
        total = sum(item['count'] for item in status_data)
        if total > 0:
            for item in status_data:
                item['percentage'] = round((item['count'] / total) * 100, 1)
        
        return jsonify({'status_distribution': status_data})
    except Exception as e:
        print(f"Error in get_status_distribution: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/recent-documents', methods=['GET'])
def get_recent_documents():
    """Get 5 most recent documents"""
    try:
        recent_docs = list(
            documents_collection.find({})
            .sort('date', -1)
            .limit(5)
        )
        recent_docs = [serialize_document(doc) for doc in recent_docs]
        
        return jsonify({'recent_documents': recent_docs})
    except Exception as e:
        print(f"Error in get_recent_documents: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/dashboard/language-distribution', methods=['GET'])
def get_language_distribution():
    """Get document distribution by language"""
    try:
        pipeline = [
            {'$group': {'_id': '$language', 'count': {'$sum': 1}}},
            {'$sort': {'count': -1}}
        ]
        
        languages = list(documents_collection.aggregate(pipeline))
        
        language_data = [
            {'language': item['_id'] or 'Unknown', 'count': item['count']}
            for item in languages
        ]
        
        return jsonify({'language_distribution': language_data})
    except Exception as e:
        print(f"Error in get_language_distribution: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>/download', methods=['GET'])
def download_document(doc_id):
    """Download a document as a text file with metadata."""
    try:
        doc = documents_collection.find_one({'_id': ObjectId(doc_id)})
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        
        # Create a formatted text file content
        content = []
        content.append(f"{'='*80}")
        content.append(f"DOCUMENT: {doc.get('title', 'Untitled')}")
        content.append(f"{'='*80}\n")
        
        # Add metadata
        content.append("METADATA:")
        content.append(f"  Department: {doc.get('department', 'N/A')}")
        content.append(f"  Type: {doc.get('type', 'N/A')}")
        content.append(f"  Date: {doc.get('date', 'N/A')}")
        content.append(f"  Status: {doc.get('status', 'N/A')}")
        content.append(f"  Source: {doc.get('source', 'N/A')}")
        content.append(f"  Language: {doc.get('language', 'N/A')}")
        content.append(f"  ID: {str(doc['_id'])}\n")
        
        # Add tags
        if doc.get('tags'):
            content.append(f"Tags: {', '.join(doc.get('tags', []))}\n")
        
        # Add summary
        if doc.get('summary'):
            content.append("SUMMARY:")
            content.append(f"{doc.get('summary', '')}\n")
        
        # Add tables
        if doc.get('tables_data') and len(doc.get('tables_data', [])) > 0:
            content.append("\nTABLES:")
            content.append("-" * 80)
            for idx, table_obj in enumerate(doc.get('tables_data', []), 1):
                caption = table_obj.get('caption', f'Table {idx}')
                content.append(f"\n{caption}:")
                
                if table_obj.get('data') and len(table_obj.get('data', [])) > 0:
                    table_data = table_obj.get('data', [])
                    # Add header
                    header = table_data[0]
                    content.append(" | ".join(str(h) for h in header))
                    content.append("-" * 80)
                    # Add rows
                    for row in table_data[1:]:
                        content.append(" | ".join(str(cell) for cell in row))
                content.append("")
        
        # Add full content
        if doc.get('content'):
            content.append("\nFULL CONTENT:")
            content.append("-" * 80)
            content.append(doc.get('content', ''))
        
        content.append(f"\n{'='*80}")
        content.append("End of Document")
        content.append(f"{'='*80}")
        
        file_content = "\n".join(content)
        
        # Create BytesIO object
        file_obj = BytesIO(file_content.encode('utf-8'))
        
        # Create sanitized filename
        safe_title = doc.get('title', 'document').replace(" ", "_").replace("/", "_")
        filename = f"{safe_title}_{str(doc['_id'])[:8]}.txt"
        
        return send_file(
            file_obj,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Error in download_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/export/csv', methods=['GET'])
def export_documents_csv():
    """Export all documents as CSV"""
    try:
        import csv
        
        documents = list(documents_collection.find({}))
        
        if not documents:
            return jsonify({'error': 'No documents to export'}), 404
        
        output = BytesIO()
        writer = csv.DictWriter(output, fieldnames=['title', 'department', 'type', 'status', 'date', 'language'])
        writer.writeheader()
        
        for doc in documents:
            writer.writerow({
                'title': doc.get('title', ''),
                'department': doc.get('department', ''),
                'type': doc.get('type', ''),
                'status': doc.get('status', ''),
                'date': doc.get('date', ''),
                'language': doc.get('language', '')
            })
        
        output.seek(0)
        
        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name='documents_export.csv'
        )
    except Exception as e:
        print(f"Error in export_documents_csv: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/bulk-update-status', methods=['PUT'])
def bulk_update_status():
    """Bulk update document status"""
    try:
        data = request.get_json()
        doc_ids = data.get('doc_ids', [])
        new_status = data.get('status', '')
        
        if not doc_ids or not new_status:
            return jsonify({'error': 'doc_ids and status are required'}), 400
        
        object_ids = [ObjectId(doc_id) for doc_id in doc_ids]
        
        result = documents_collection.update_many(
            {'_id': {'$in': object_ids}},
            {'$set': {'status': new_status}}
        )
        
        return jsonify({
            'message': f'Updated {result.modified_count} documents',
            'modified_count': result.modified_count
        })
    except Exception as e:
        print(f"Error in bulk_update_status: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/search-advanced', methods=['POST'])
def advanced_search():
    """Advanced search with multiple filters"""
    try:
        data = request.get_json()
        search_query = data.get('search', '')
        department = data.get('department', '')
        doc_type = data.get('type', '')
        status = data.get('status', '')
        date_from = data.get('date_from', '')
        date_to = data.get('date_to', '')
        tags = data.get('tags', [])
        
        query_filter = {}
        
        if search_query:
            query_filter['$or'] = [
                {'title': {'$regex': search_query, '$options': 'i'}},
                {'summary': {'$regex': search_query, '$options': 'i'}},
                {'tags': {'$regex': search_query, '$options': 'i'}},
                {'content': {'$regex': search_query, '$options': 'i'}}
            ]
        
        if department:
            query_filter['department'] = department
        
        if doc_type:
            query_filter['type'] = doc_type.replace('-', ' ').title()
        
        if status:
            query_filter['status'] = status
        
        if date_from or date_to:
            date_range = {}
            if date_from:
                date_range['$gte'] = datetime.strptime(date_from, '%Y-%m-%d')
            if date_to:
                date_range['$lte'] = datetime.strptime(date_to, '%Y-%m-%d')
            if date_range:
                query_filter['date'] = date_range
        
        if tags and len(tags) > 0:
            query_filter['tags'] = {'$in': tags}
        
        documents = list(documents_collection.find(query_filter).sort('date', -1))
        documents = [serialize_document(doc) for doc in documents]
        
        return jsonify({
            'results': documents,
            'count': len(documents)
        })
    except Exception as e:
        print(f"Error in advanced_search: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/<doc_id>/star', methods=['POST'])
def star_document(doc_id):
    """Star/unstar a document"""
    try:
        doc = documents_collection.find_one({'_id': ObjectId(doc_id)})
        if not doc:
            return jsonify({'error': 'Document not found'}), 404
        
        new_starred_state = not doc.get('starred', False)
        
        documents_collection.update_one(
            {'_id': ObjectId(doc_id)},
            {'$set': {'starred': new_starred_state}}
        )
        
        return jsonify({
            'message': f'Document {"starred" if new_starred_state else "unstarred"} successfully',
            'starred': new_starred_state
        })
    except Exception as e:
        print(f"Error in star_document: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        # Try to connect to database
        documents_collection.count_documents({})
        return jsonify({'status': 'healthy', 'database': 'connected'})
    except Exception as e:
        return jsonify({'status': 'unhealthy', 'error': str(e)}), 500


@app.route('/api/documents/starred', methods=['GET'])
def get_starred_documents():
    """Get all starred documents"""
    try:
        starred_docs = list(
            documents_collection.find({'starred': True})
            .sort('date', -1)
        )
        starred_docs = [serialize_document(doc) for doc in starred_docs]
        
        return jsonify({
            'documents': starred_docs,
            'count': len(starred_docs)
        })
    except Exception as e:
        print(f"Error in get_starred_documents: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/documents/filter-by-department/<department>', methods=['GET'])
def get_documents_by_department(department):
    """Get all documents from a specific department"""
    try:
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 10))
        
        total = documents_collection.count_documents({'department': department})
        
        skip = (page - 1) * limit
        documents = list(
            documents_collection.find({'department': department})
            .sort('date', -1)
            .skip(skip)
            .limit(limit)
        )
        documents = [serialize_document(doc) for doc in documents]
        
        return jsonify({
            'documents': documents,
            'pagination': {
                'total': total,
                'page': page,
                'limit': limit,
                'pages': (total + limit - 1) // limit
            }
        })
    except Exception as e:
        print(f"Error in get_documents_by_department: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    return jsonify({'error': 'Endpoint not found'}), 404


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors"""
    return jsonify({'error': 'Internal server error'}), 500


if __name__ == '__main__':
    if not GEMINI_API_KEY:
        print("🚨 Warning: GEMINI_API_KEY environment variable is not set.")
        print("AI features will be disabled, and mock data will be used.")
    app.run(debug=True, port=5000)